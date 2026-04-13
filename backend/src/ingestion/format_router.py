"""Format Router: detect file type and dispatch to the appropriate parser."""

from __future__ import annotations

import logging
from pathlib import Path

from src.api.models import ParsedDocument

logger = logging.getLogger(__name__)


class FormatRouter:
    """Parse PDF / DOCX / HTML / plain-text files into raw text."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".txt"}

    async def parse(
        self,
        file_path: str,
        file_content: bytes | None = None,
    ) -> ParsedDocument:
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.warning("Unsupported file extension: %s — treating as plain text", ext)
            ext = ".txt"

        if ext == ".pdf":
            raw_text, hints = self._parse_pdf(file_path)
        elif ext == ".docx":
            raw_text, hints = self._parse_docx(file_path)
        elif ext in {".html", ".htm"}:
            raw_text, hints = self._parse_html(file_path, file_content)
        else:
            raw_text, hints = self._parse_text(file_path, file_content)

        return ParsedDocument(
            raw_text=raw_text,
            format_hints=hints,
            file_path=file_path,
        )

    # ------------------------------------------------------------------
    # PDF — PyMuPDF (fitz)
    # ------------------------------------------------------------------
    def _parse_pdf(self, file_path: str) -> tuple[str, dict]:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        pages: list[str] = []
        for page in doc:
            text = page.get_text("text")
            pages.append(text)
        doc.close()

        raw_text = "\n\n".join(pages)

        # Detect if scanned PDF (very little text extracted)
        char_count = len(raw_text.strip())
        hints: dict = {"format": "pdf", "pages": len(pages)}
        if char_count < 100 and len(pages) > 0:
            logger.warning(
                "PDF %s has very little text (%d chars) — likely scanned. "
                "Consider using LlamaParse for OCR.",
                file_path,
                char_count,
            )
            hints["likely_scanned"] = True

        return raw_text, hints

    # ------------------------------------------------------------------
    # DOCX — python-docx
    # ------------------------------------------------------------------
    def _parse_docx(self, file_path: str) -> tuple[str, dict]:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(file_path)
        parts: list[str] = []
        headings_found: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()

            # Preserve heading information as hints
            if "heading" in style_name:
                level = style_name.replace("heading", "").strip() or "1"
                parts.append(f"[H{level}] {text}")
                headings_found.append(text)
            elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                # Centered text in Vietnamese legal docs is often a title/header
                parts.append(f"[CENTER] {text}")
            else:
                parts.append(text)

        # Also extract text from tables
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                parts.append("[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]")

        raw_text = "\n".join(parts)
        hints: dict = {
            "format": "docx",
            "headings_found": headings_found,
            "tables_found": len(doc.tables),
        }
        return raw_text, hints

    # ------------------------------------------------------------------
    # HTML — BeautifulSoup
    # ------------------------------------------------------------------
    def _parse_html(
        self, file_path: str, file_content: bytes | None = None
    ) -> tuple[str, dict]:
        from bs4 import BeautifulSoup

        if file_content:
            html = file_content.decode("utf-8", errors="replace")
        else:
            with open(file_path, encoding="utf-8", errors="replace") as f:
                html = f.read()

        soup = BeautifulSoup(html, "lxml")

        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        parts: list[str] = []
        headings_found: list[str] = []

        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "div"]
        ):
            text = element.get_text(strip=True)
            if not text:
                continue

            tag_name = element.name
            if tag_name and tag_name.startswith("h"):
                level = tag_name[1]
                parts.append(f"[H{level}] {text}")
                headings_found.append(text)
            elif tag_name == "li":
                parts.append(f"- {text}")
            else:
                parts.append(text)

        raw_text = "\n".join(parts)
        hints: dict = {
            "format": "html",
            "headings_found": headings_found,
        }
        return raw_text, hints

    # ------------------------------------------------------------------
    # Plain text
    # ------------------------------------------------------------------
    def _parse_text(
        self, file_path: str, file_content: bytes | None = None
    ) -> tuple[str, dict]:
        if file_content:
            raw_text = file_content.decode("utf-8", errors="replace")
        else:
            with open(file_path, encoding="utf-8", errors="replace") as f:
                raw_text = f.read()
        return raw_text, {"format": "text"}

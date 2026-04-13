"""Export rendered contract content to DOCX and PDF formats."""

from __future__ import annotations

import io
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


def _is_header_line(line: str) -> bool:
    """Check if a line looks like a section header (all caps or starts with DIEU/BEN)."""
    stripped = line.strip()
    if not stripped:
        return False
    return (
        stripped.isupper()
        or stripped.startswith("DIEU ")
        or stripped.startswith("BEN ")
        or stripped.startswith("QUYET DINH")
        or stripped.startswith("NOI DUNG")
        or stripped.startswith("KET LUAN")
        or stripped.startswith("LUU Y VE")
    )


def _is_centered_line(line: str) -> bool:
    """Check if a line should be centered (title block lines)."""
    stripped = line.strip()
    return stripped in (
        "CONG HOA XA HOI CHU NGHIA VIET NAM",
        "Doc lap — Tu do — Hanh phuc",
        "---------",
        "HOP DONG LAO DONG",
        "HOP DONG THU VIEC",
        "HOP DONG THUE MAT BANG",
        "HOP DONG DICH VU",
        "QUYET DINH",
        "BIEN BAN VI PHAM KY LUAT LAO DONG",
        "(Xac dinh thoi han)",
        "(Khong xac dinh thoi han)",
    ) or stripped.startswith("So: ")


def to_docx(rendered_content: str, title: str = "") -> bytes:
    """Convert rendered contract text to DOCX format.

    Args:
        rendered_content: The rendered contract text from template engine.
        title: Optional document title for metadata.

    Returns:
        DOCX file as bytes.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    if title:
        doc.core_properties.title = title

    lines = rendered_content.split("\n")
    in_title_block = True

    for line in lines:
        stripped = line.strip()

        # Skip pure separator lines
        if stripped == "---------":
            continue

        # Empty line = paragraph break
        if not stripped:
            doc.add_paragraph("")
            in_title_block = False
            continue

        # Centered title block (first few lines)
        if in_title_block and _is_centered_line(stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            if stripped.isupper() or stripped.startswith("HOP DONG") or stripped == "QUYET DINH":
                run.bold = True
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            continue

        in_title_block = False

        # Section headers (DIEU X, BEN A, etc.)
        if _is_header_line(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Signature block detection
        if "BEN A" in stripped and "BEN B" in stripped:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Sub-items (indented with a), b), etc.)
        if re.match(r"^\s*[a-z]\)", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def to_pdf(rendered_content: str, title: str = "") -> bytes:
    """Convert rendered contract text to PDF format via WeasyPrint.

    Args:
        rendered_content: The rendered contract text from template engine.
        title: Optional document title.

    Returns:
        PDF file as bytes.
    """
    from weasyprint import HTML

    # Escape HTML special chars in content
    escaped = (
        rendered_content
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )

    # Convert newlines to HTML paragraphs
    lines = escaped.split("\n")
    html_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            html_lines.append("<br/>")
        elif stripped == "---------":
            html_lines.append('<hr style="border: none; border-top: 1px solid #000;"/>')
        elif _is_centered_line(stripped.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")):
            if stripped.isupper() or "HOP DONG" in stripped or stripped == "QUYET DINH":
                html_lines.append(f'<p style="text-align:center;font-weight:bold;font-size:14pt;">{stripped}</p>')
            else:
                html_lines.append(f'<p style="text-align:center;">{stripped}</p>')
        elif _is_header_line(stripped.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")):
            html_lines.append(f"<p><strong>{stripped}</strong></p>")
        else:
            html_lines.append(f"<p>{stripped}</p>")

    body = "\n".join(html_lines)

    html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<title>{title or "Hop dong"}</title>
<style>
@page {{
    size: A4;
    margin: 2cm 2cm 2cm 2.5cm;
}}
body {{
    font-family: "Times New Roman", Times, serif;
    font-size: 12pt;
    line-height: 1.5;
    color: #000;
}}
p {{
    margin: 2pt 0;
}}
</style>
</head>
<body>
{body}
</body>
</html>"""

    pdf_bytes = HTML(string=html_content).write_pdf()
    return pdf_bytes
